"""
Quick Text Extraction Service

Provides fast text extraction from PDF and DOCX files without full document processing.
Used for immediate AI responses when users attach files to chat.
"""

import logging
import io
import uuid
from typing import Dict, Any, Optional, List, Tuple
import tempfile
import os

logger = logging.getLogger(__name__)

# Maximum pages to extract for quick mode (to prevent memory issues)
MAX_QUICK_EXTRACT_PAGES = 50


def detect_file_type(file_bytes: bytes, filename: str) -> str:
    """
    Detect file type from bytes and filename.
    
    Args:
        file_bytes: Raw file content
        filename: Original filename
        
    Returns:
        File type string: 'pdf', 'docx', 'doc', 'txt', or 'unknown'
    """
    filename_lower = filename.lower() if filename else ''
    
    # Check by extension first
    if filename_lower.endswith('.pdf'):
        return 'pdf'
    elif filename_lower.endswith('.docx'):
        return 'docx'
    elif filename_lower.endswith('.doc'):
        return 'doc'
    elif filename_lower.endswith('.txt'):
        return 'txt'
    
    # Check magic bytes
    if len(file_bytes) >= 4:
        # PDF magic bytes: %PDF
        if file_bytes[:4] == b'%PDF':
            return 'pdf'
        # DOCX magic bytes (ZIP with specific structure)
        if file_bytes[:4] == b'PK\x03\x04':
            return 'docx'
    
    return 'unknown'


def extract_text_from_pdf(file_bytes: bytes, max_pages: int = MAX_QUICK_EXTRACT_PAGES) -> Dict[str, Any]:
    """
    Extract text from PDF using PyMuPDF (fitz).
    
    Args:
        file_bytes: Raw PDF content
        max_pages: Maximum number of pages to extract
        
    Returns:
        Dictionary with:
        - success: bool
        - text: Full concatenated text
        - page_texts: List of text per page
        - page_count: Total pages in document
        - extracted_pages: Number of pages actually extracted
        - error: Error message if failed
    """
    try:
        import fitz  # type: ignore # PyMuPDF - imported as 'fitz'
        
        # Open PDF from bytes
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        
        total_pages = len(pdf_document)
        pages_to_extract = min(total_pages, max_pages)
        
        page_texts: List[str] = []
        
        for page_num in range(pages_to_extract):
            page = pdf_document[page_num]
            text = page.get_text("text")
            page_texts.append(text.strip())
        
        pdf_document.close()
        
        # Concatenate all text with page markers
        full_text = "\n\n".join([
            f"--- Page {i+1} ---\n{text}" 
            for i, text in enumerate(page_texts) 
            if text
        ])
        
        truncated = total_pages > max_pages
        
        logger.info(f"📄 Quick PDF extraction: {pages_to_extract}/{total_pages} pages, {len(full_text)} chars")
        
        return {
            'success': True,
            'text': full_text,
            'page_texts': page_texts,
            'page_count': total_pages,
            'extracted_pages': pages_to_extract,
            'truncated': truncated,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }
        
    except ImportError:
        logger.error("❌ PyMuPDF (fitz) not installed. Run: pip install pymupdf")
        return {
            'success': False,
            'error': 'PDF extraction library not available',
            'text': '',
            'page_texts': [],
            'page_count': 0
        }
    except Exception as e:
        logger.error(f"❌ PDF extraction failed: {str(e)}")
        return {
            'success': False,
            'error': str(e),
            'text': '',
            'page_texts': [],
            'page_count': 0
        }


def extract_text_from_docx(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from DOCX using python-docx.
    
    Args:
        file_bytes: Raw DOCX content
        
    Returns:
        Dictionary with:
        - success: bool
        - text: Full document text
        - page_texts: List with single entry (DOCX doesn't have true pages)
        - page_count: Estimated page count
        - error: Error message if failed
    """
    try:
        from docx import Document  # type: ignore # python-docx
        
        # Create a file-like object from bytes
        file_stream = io.BytesIO(file_bytes)
        doc = Document(file_stream)
        
        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
        
        full_text = '\n\n'.join(paragraphs)
        
        # Estimate page count (rough: ~3000 chars per page)
        estimated_pages = max(1, len(full_text) // 3000)
        
        logger.info(f"📄 Quick DOCX extraction: ~{estimated_pages} pages, {len(full_text)} chars")
        
        return {
            'success': True,
            'text': full_text,
            'page_texts': [full_text],  # DOCX doesn't have true page breaks
            'page_count': estimated_pages,
            'extracted_pages': estimated_pages,
            'truncated': False,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }
        
    except ImportError:
        logger.error("❌ python-docx not installed. Run: pip install python-docx")
        return {
            'success': False,
            'error': 'DOCX extraction library not available',
            'text': '',
            'page_texts': [],
            'page_count': 0
        }
    except Exception as e:
        logger.error(f"❌ DOCX extraction failed: {str(e)}")
        return {
            'success': False,
            'error': str(e),
            'text': '',
            'page_texts': [],
            'page_count': 0
        }


def extract_text_from_txt(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from plain text file.
    
    Args:
        file_bytes: Raw text content
        
    Returns:
        Dictionary with extracted text
    """
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = file_bytes.decode('utf-8', errors='replace')
        
        # Estimate page count
        estimated_pages = max(1, len(text) // 3000)
        
        logger.info(f"📄 Quick TXT extraction: {len(text)} chars")
        
        return {
            'success': True,
            'text': text,
            'page_texts': [text],
            'page_count': estimated_pages,
            'extracted_pages': estimated_pages,
            'truncated': False,
            'char_count': len(text),
            'word_count': len(text.split())
        }
        
    except Exception as e:
        logger.error(f"❌ TXT extraction failed: {str(e)}")
        return {
            'success': False,
            'error': str(e),
            'text': '',
            'page_texts': [],
            'page_count': 0
        }


def quick_extract(file_bytes: bytes, filename: str, store_temp: bool = True) -> Dict[str, Any]:
    """
    Main entry point for quick text extraction.
    
    Args:
        file_bytes: Raw file content
        filename: Original filename
        store_temp: Whether to store file temporarily for later full processing
        
    Returns:
        Dictionary with:
        - success: bool
        - text: Extracted text
        - page_texts: List of text per page
        - page_count: Number of pages
        - file_type: Detected file type
        - temp_file_id: UUID for temp storage (if store_temp=True)
        - error: Error message if failed
    """
    file_type = detect_file_type(file_bytes, filename)
    
    logger.info(f"🔍 Quick extract starting for {filename} (type: {file_type}, size: {len(file_bytes)} bytes)")
    
    # Extract based on file type
    if file_type == 'pdf':
        result = extract_text_from_pdf(file_bytes)
    elif file_type == 'docx':
        result = extract_text_from_docx(file_bytes)
    elif file_type == 'txt':
        result = extract_text_from_txt(file_bytes)
    else:
        return {
            'success': False,
            'error': f'Unsupported file type: {file_type}. Supported: PDF, DOCX, TXT',
            'text': '',
            'page_texts': [],
            'page_count': 0,
            'file_type': file_type
        }
    
    result['file_type'] = file_type
    result['filename'] = filename
    
    # Generate temp file ID for potential later full processing
    if store_temp and result['success']:
        result['temp_file_id'] = str(uuid.uuid4())
    
    return result


def store_temp_file(file_bytes: bytes, filename: str, temp_file_id: str) -> Dict[str, Any]:
    """
    Store file temporarily in S3 for later full processing.
    
    Args:
        file_bytes: Raw file content
        filename: Original filename
        temp_file_id: UUID for temp storage
        
    Returns:
        Dictionary with storage result
    """
    try:
        import boto3
        import os
        
        bucket_name = os.environ.get('S3_UPLOAD_BUCKET')
        if not bucket_name:
            logger.warning("⚠️ S3_UPLOAD_BUCKET not configured, temp storage disabled")
            return {'success': False, 'error': 'S3 not configured'}
        
        s3_client = boto3.client('s3')
        
        # Store with temp_ prefix for easy cleanup
        s3_key = f"temp_uploads/{temp_file_id}/{filename}"
        
        s3_client.put_object(
            Bucket=bucket_name,
            Key=s3_key,
            Body=file_bytes,
            ContentType='application/octet-stream',
            Metadata={
                'temp_file_id': temp_file_id,
                'original_filename': filename
            }
        )
        
        logger.info(f"✅ Stored temp file: {s3_key}")
        
        return {
            'success': True,
            's3_key': s3_key,
            'temp_file_id': temp_file_id
        }
        
    except Exception as e:
        logger.error(f"❌ Failed to store temp file: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def get_temp_file(temp_file_id: str) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Retrieve a temporarily stored file.
    
    Args:
        temp_file_id: UUID of temp file
        
    Returns:
        Tuple of (file_bytes, filename) or (None, None) if not found
    """
    try:
        import boto3
        import os
        
        bucket_name = os.environ.get('S3_UPLOAD_BUCKET')
        if not bucket_name:
            return None, None
        
        s3_client = boto3.client('s3')
        
        # List objects with temp prefix
        prefix = f"temp_uploads/{temp_file_id}/"
        response = s3_client.list_objects_v2(Bucket=bucket_name, Prefix=prefix)
        
        if 'Contents' not in response or len(response['Contents']) == 0:
            return None, None
        
        # Get the first (and should be only) file
        s3_key = response['Contents'][0]['Key']
        filename = s3_key.split('/')[-1]
        
        obj = s3_client.get_object(Bucket=bucket_name, Key=s3_key)
        file_bytes = obj['Body'].read()
        
        return file_bytes, filename
        
    except Exception as e:
        logger.error(f"❌ Failed to retrieve temp file: {str(e)}")
        return None, None


def delete_temp_file(temp_file_id: str) -> bool:
    """
    Delete a temporarily stored file.
    
    Args:
        temp_file_id: UUID of temp file
        
    Returns:
        True if deleted, False otherwise
    """
    try:
        import boto3
        import os
        
        bucket_name = os.environ.get('S3_UPLOAD_BUCKET')
        if not bucket_name:
            return False
        
        s3_client = boto3.client('s3')
        
        # List and delete objects with temp prefix
        prefix = f"temp_uploads/{temp_file_id}/"
        response = s3_client.list_objects_v2(Bucket=bucket_name, Prefix=prefix)
        
        if 'Contents' in response:
            for obj in response['Contents']:
                s3_client.delete_object(Bucket=bucket_name, Key=obj['Key'])
            logger.info(f"🗑️ Deleted temp file: {temp_file_id}")
            return True
        
        return False
        
    except Exception as e:
        logger.error(f"❌ Failed to delete temp file: {str(e)}")
        return False

